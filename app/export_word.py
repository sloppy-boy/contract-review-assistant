"""报告 → Word 导出（python-docx，SPEC 2.8 交付物）。

Web 为主 + 文档落地的行业标准组合：浏览器在线审查 → 正式 Word 审阅报告。
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SEV_COLOR = {"high": RGBColor(0xDC, 0x26, 0x26), "medium": RGBColor(0xD9, 0x77, 0x06), "low": RGBColor(0x05, 0x96, 0x69)}
SEV_NAME = {"high": "高危", "medium": "中危", "low": "低危"}
TIER_NAME = {"direct": "直接依据", "indirect": "间接依据", "none": "无直接依据"}


def _set_cn_font(run, name: str = "微软雅黑", size: float = 10.5) -> None:
    """设置中西文字体（python-docx 需显式指定 East Asian 字体，否则中文走默认）。"""
    run.font.name = name
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def _add_para(doc: Document, text: str, size: float = 10.5, bold: bool = False, color=None, align=None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_cn_font(run, size=size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p


def report_to_docx(report: dict) -> bytes:
    """报告 dict → .docx 字节。结构对齐统一 JSON schema（评测/前端/Word 三处共用）。"""
    doc = Document()
    # 页边距收窄，多放内容
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)

    _add_para(doc, "合同审查报告", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "多 Agent 协同审查 · 法条可溯源 · 人工终审", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    contract = report.get("contract", {})
    meta = report.get("meta", {})
    _add_para(doc, f"合同：{contract.get('name', '—')}　类型：{'采购' if contract.get('type') == 'purchase' else '销售'}"
                   f"　条款数：{contract.get('clauseCount', 0)}", size=10.5, bold=True)
    _add_para(doc, f"生成时间：{meta.get('generatedAt', '—')}　复核模式：{meta.get('reviewMode', '—')}"
                   f"　{'演示数据（mock）' if meta.get('mock') else '真实 pipeline 输出'}", size=9)

    s = report.get("summary", {})
    doc.add_paragraph()
    _add_para(doc, f"风险统计：🔴 高危 {s.get('high', 0)}　🟡 中危 {s.get('medium', 0)}　🟢 低危 {s.get('low', 0)}　合计 {s.get('total', 0)}",
              size=12, bold=True)
    doc.add_paragraph()

    # ---- 风险清单 ----
    risks = report.get("risks", [])
    if not risks:
        _add_para(doc, "未发现风险条款。", size=11)
    for i, r in enumerate(risks, 1):
        sev = r.get("severity", "low")
        color = SEV_COLOR.get(sev, RGBColor(0x4B, 0x55, 0x63))
        _add_para(doc, f"{i}. 【{SEV_NAME.get(sev, sev)}】{r.get('riskType', '')}"
                       f"　（条款 {r.get('clauseId', '')}）"
                       + ("　【有争议】" if r.get("disputed") else ""),
                  size=12, bold=True, color=color)
        _add_para(doc, f"原文摘录：{r.get('clauseQuote', '')}", size=10)
        lb = r.get("legalBasis", {})
        tier = lb.get("tier", "none")
        if tier == "direct":
            _add_para(doc, f"法条依据【直接依据】：{lb.get('articleId', '')} · {lb.get('version', '')}", size=10)
            _add_para(doc, f"　　条文原文：{lb.get('quote', '')}", size=9)
        elif tier == "indirect":
            _add_para(doc, "法条依据【间接依据】：诚信/公平原则等原则性条款（非直接条文）", size=10)
        else:
            _add_para(doc, "法条依据【无直接依据】：提示性质，无明确法条依据（未硬编）", size=10)
        _add_para(doc, f"证据：{r.get('evidence', '')}", size=10)
        _add_para(doc, f"修改建议：{r.get('suggestion', '')}", size=10)
        if r.get("suggestionClauseText"):
            _add_para(doc, f"示范条款：{r.get('suggestionClauseText', '')}", size=9.5, color=RGBColor(0x1D, 0x4E, 0xD8))
        if r.get("status") == "disputed" and r.get("reVerifyJustification"):
            _add_para(doc, f"（复核驳回后 worker 坚持：{r.get('reVerifyJustification')}）", size=9)
        doc.add_paragraph()

    # ---- 页脚声明 ----
    doc.add_paragraph()
    _add_para(doc, "— 本报告由 AI 初筛生成，输出需经人工终审，不构成法律意见 —", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
